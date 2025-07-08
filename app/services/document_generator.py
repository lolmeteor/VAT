"""
Сервис для генерации документов (TXT, DOCX) из данных БД
"""
from docx import Document
from io import BytesIO
import json

from app.models import Transcription, Analysis

class DocumentGeneratorService:
    """
    Сервис для генерации документов на лету.
    """

    def generate_transcription_txt(self, transcription: Transcription) -> bytes:
        """Генерирует TXT файл с транскрипцией."""
        if not transcription.transcription_text:
            return "Текст транскрипции отсутствует.".encode('utf-8')

        header = (
            f"ТРАНСКРИПЦИЯ АУДИОФАЙЛА\n"
            f"=========================\n"
            f"Дата: {transcription.created_at.strftime('%d.%m.%Y %H:%M')}\n"
            f"Язык: {transcription.language_detected or 'Не определен'}\n"
            f"Количество спикеров: {transcription.speakers_count or 'Не определено'}\n"
            f"=========================\n\n"
        )
        
        content = header + transcription.transcription_text
        return content.encode('utf-8')

    def generate_transcription_docx(self, transcription: Transcription) -> bytes:
        """Генерирует DOCX файл с ранскрипцией."""
        doc = Document()
        doc.add_heading('Транскрипция аудиофайла', 0)

        p = doc.add_paragraph()
        p.add_run('Дата: ').bold = True
        p.add_run(transcription.created_at.strftime('%d.%m.%Y %H:%M'))

        p = doc.add_paragraph()
        p.add_run('Язык: ').bold = True
        p.add_run(transcription.language_detected or 'Не определен')
        
        p = doc.add_paragraph()
        p.add_run('Количество спикеров: ').bold = True
        p.add_run(str(transcription.speakers_count or 'Не определено'))
        
        doc.add_paragraph('') # Пустая строка для отступа
        
        doc.add_heading('Текст транскрипции', level=1)
        doc.add_paragraph(transcription.transcription_text or "Текст отсутствует.")

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()

    def generate_analysis_docx(self, analysis: Analysis) -> bytes:
        """Генерирует DOCX файл с анализом."""
        doc = Document()
        doc.add_heading(f'Отчет по анализу: "{analysis.analysis_type.value}"', 0)

        if analysis.analysis_summary:
            doc.add_heading('Краткое резюме', level=1)
            doc.add_paragraph(analysis.analysis_summary)

        if analysis.key_points:
            doc.add_heading('Ключевые моменты', level=1)
            try:
                # Обрабатываем JSON из БД
                points = json.loads(analysis.key_points) if isinstance(analysis.key_points, str) else analysis.key_points
                if isinstance(points, list):
                    for point in points:
                        doc.add_paragraph(str(point), style='List Bullet')
                else:
                    doc.add_paragraph(str(points))
            except (json.JSONDecodeError, TypeError):
                doc.add_paragraph("Не удалось обработать ключевые моменты.")

        if analysis.analysis_text:
            doc.add_heading('Подробный отчет', level=1)
            doc.add_paragraph(analysis.analysis_text)

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()
